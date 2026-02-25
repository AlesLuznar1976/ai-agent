"""
File Processor - Procesira naložene datoteke za Claude Vision API.

Podpira:
- Slike (image/*) → base64 → Claude image content block
- PDF (.pdf) → base64 → Claude document content block
- Excel (.xlsx) → openpyxl → tekst
- Word (.docx) → python-docx → tekst
- CSV (.csv) → prebere kot tekst
- Ostalo → vrne filename + info
"""

import base64
import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

IMAGE_MIME_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}
PDF_MIME_TYPE = "application/pdf"


def process_uploaded_file(filepath: str, mime_type: str) -> dict:
    """
    Procesira naloženo datoteko in vrne Claude content block.

    Returns:
        dict z ključi:
        - type: "image" | "document" | "text" | "unsupported"
        - content_block: Claude API content block (za image/document)
        - text: ekstrahiran tekst (za Excel/Word/CSV)
        - filename: ime datoteke
    """
    path = Path(filepath)
    filename = path.name
    suffix = path.suffix.lower()

    # Slike → Claude image content block
    if mime_type in IMAGE_MIME_TYPES:
        return _process_image(filepath, mime_type, filename)

    # PDF → Claude document content block
    if mime_type == PDF_MIME_TYPE or suffix == ".pdf":
        return _process_pdf(filepath, filename)

    # Excel
    if suffix in (".xlsx", ".xls") or mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _process_excel(filepath, filename)

    # Word
    if suffix in (".docx", ".doc") or mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _process_word(filepath, filename)

    # CSV
    if suffix == ".csv" or mime_type == "text/csv":
        return _process_csv(filepath, filename)

    # Tekst
    if mime_type and mime_type.startswith("text/"):
        return _process_text(filepath, filename)

    # Nepoznan format
    return {
        "type": "unsupported",
        "filename": filename,
        "text": f"[Datoteka: {filename} ({mime_type or 'neznan tip'})]",
    }


def _process_image(filepath: str, mime_type: str, filename: str) -> dict:
    """Slika → base64 → Claude image content block."""
    with open(filepath, "rb") as f:
        data = base64.standard_b64encode(f.read()).decode("utf-8")

    return {
        "type": "image",
        "filename": filename,
        "content_block": {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": mime_type,
                "data": data,
            },
        },
    }


def _process_pdf(filepath: str, filename: str) -> dict:
    """PDF → base64 → Claude document content block."""
    with open(filepath, "rb") as f:
        data = base64.standard_b64encode(f.read()).decode("utf-8")

    return {
        "type": "document",
        "filename": filename,
        "content_block": {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": data,
            },
        },
    }


def _process_excel(filepath: str, filename: str) -> dict:
    """Excel → openpyxl → tekst."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(filepath, read_only=True, data_only=True)
        lines = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== List: {sheet_name} ===")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                lines.append("\t".join(cells))
                row_count += 1
                if row_count >= 500:
                    lines.append(f"... (prikazanih prvih 500 od več vrstic)")
                    break
            lines.append("")

        wb.close()
        text = "\n".join(lines)

        return {
            "type": "text",
            "filename": filename,
            "text": f"[Vsebina Excel datoteke: {filename}]\n\n{text}",
        }

    except Exception as e:
        logger.error(f"Napaka pri branju Excel datoteke {filename}: {e}")
        return {
            "type": "text",
            "filename": filename,
            "text": f"[Napaka pri branju Excel datoteke {filename}: {e}]",
        }


def _process_word(filepath: str, filename: str) -> dict:
    """Word → python-docx → tekst."""
    try:
        from docx import Document

        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        # Tudi tabele
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("\t".join(cells))
            if rows:
                text += "\n\n[Tabela]\n" + "\n".join(rows)

        return {
            "type": "text",
            "filename": filename,
            "text": f"[Vsebina Word dokumenta: {filename}]\n\n{text}",
        }

    except Exception as e:
        logger.error(f"Napaka pri branju Word datoteke {filename}: {e}")
        return {
            "type": "text",
            "filename": filename,
            "text": f"[Napaka pri branju Word datoteke {filename}: {e}]",
        }


def _process_csv(filepath: str, filename: str) -> dict:
    """CSV → tekst."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            lines = []
            for i, row in enumerate(reader):
                lines.append("\t".join(row))
                if i >= 500:
                    lines.append("... (prikazanih prvih 500 vrstic)")
                    break

        text = "\n".join(lines)

        return {
            "type": "text",
            "filename": filename,
            "text": f"[Vsebina CSV datoteke: {filename}]\n\n{text}",
        }

    except Exception as e:
        logger.error(f"Napaka pri branju CSV datoteke {filename}: {e}")
        return {
            "type": "text",
            "filename": filename,
            "text": f"[Napaka pri branju CSV datoteke {filename}: {e}]",
        }


def _process_text(filepath: str, filename: str) -> dict:
    """Tekstovna datoteka → tekst."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read(100_000)  # Max 100KB teksta

        return {
            "type": "text",
            "filename": filename,
            "text": f"[Vsebina datoteke: {filename}]\n\n{text}",
        }

    except Exception as e:
        logger.error(f"Napaka pri branju tekstovne datoteke {filename}: {e}")
        return {
            "type": "text",
            "filename": filename,
            "text": f"[Napaka pri branju datoteke {filename}: {e}]",
        }

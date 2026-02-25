"""
Document Templates - Profesionalne Word predloge za Luznar d.o.o.

Podpira 4 tipe dokumentov:
- Reklamacija (Supplier Quality Complaint - po vzoru 100100306.pdf)
- RFQ analiza (analiza povpraševanja)
- BOM pregled (pregled kosovnice)
- Poročilo o pregledu (splošno poročilo/zapisnik)

Vsak template dobi strukturirane podatke (dict) in generira
profesionalen .docx z Luznar brandingom.
"""

import io
import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


# Luznar barve
NAVY = RGBColor(0x0A, 0x1A, 0x2F)
NAVY_HEX = "0A1A2F"
GOLD = RGBColor(0xB8, 0x96, 0x3E)
GOLD_HEX = "B8963E"
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_HEX = "F3F4F6"
RED = RGBColor(0xDC, 0x26, 0x26)


DOCUMENT_TYPES = {
    "reklamacija": {
        "title": "SUPPLIER QUALITY COMPLAINT",
        "subtitle": "Quality Non-Conformance Report",
    },
    "rfq_analiza": {
        "title": "Analiza povpraševanja (RFQ)",
        "subtitle": "Luznar d.o.o. - Elektronske vezje",
    },
    "bom_pregled": {
        "title": "Pregled kosovnice (BOM)",
        "subtitle": "Luznar d.o.o. - Elektronske vezje",
    },
    "porocilo": {
        "title": "Poročilo o pregledu",
        "subtitle": "Luznar d.o.o. - Elektronske vezje",
    },
}

# Claude prompt za ekstrakcijo strukturiranih podatkov
EXTRACTION_PROMPTS = {
    "reklamacija": """Iz spodnje analize/dokumenta izvleci podatke za SUPPLIER QUALITY COMPLAINT (reklamacijo dobavitelju).
Vrni SAMO JSON (brez markdown, brez ```), s točno temi ključi:

{
  "complaint_reference": "številka reklamacije (npr. 100100XXX)",
  "complaint_date": "DD Month YYYY (v angleščini, npr. 24 February 2026)",
  "supplier_company": "ime dobavitelja",
  "supplier_address": "naslov dobavitelja",
  "supplier_contact": "kontaktna oseba dobavitelja",
  "customer_company": "LUZNAR electronics d.o.o.",
  "customer_address": "Hrastje 52g, SI-4000 Kranj, Slovenia",
  "customer_contact": "kontakt email",
  "article": "artikel koda - naziv artikla",
  "purchase_order": "številka naročila",
  "delivery_date": "DD Month YYYY",
  "quantity_delivered": "količina dostavljena (npr. 100 pcs)",
  "quantity_rejected": "količina zavrnjena (npr. 100 pcs (100%))",
  "detection_point": "točka zaznave (npr. Incoming Quality Inspection)",
  "non_conformance_type": "tip neskladnosti (npr. Inadequate product quality - Missing mechanical features)",
  "description": "podroben opis napake (več stavkov, opisno)",
  "corrective_actions": ["seznam zahtevanih korektivnih ukrepov - vsak kot en stavek"],
  "supplier_decision": "odločitev: Return to supplier / Rework at supplier / Replacement quantity (če ni podatka, pusti prazno)",
  "deadline_4d": "rok za 4D poročilo (DD Month YYYY) - če ni podatka, izračunaj 7 dni od complaint_date",
  "deadline_8d": "rok za 8D poročilo (DD Month YYYY) - če ni podatka, izračunaj 28 dni od complaint_date",
  "photo_descriptions": ["opisi fotografij če obstajajo"]
}""",

    "rfq_analiza": """Iz spodnje analize/dokumenta izvleci podatke za ANALIZO RFQ POVPRAŠEVANJA.
Vrni SAMO JSON (brez markdown, brez ```), s točno temi ključi:

{
  "stevilka_rfq": "RFQ številka če je razvidna, sicer ''",
  "datum_prejema": "YYYY-MM-DD",
  "stranka": "ime stranke",
  "kontakt": "kontaktna oseba če je razvidna",
  "opis_projekta": "kratek opis projekta/izdelka (2-3 stavki)",
  "tip_vezja": "tip PCB (enostransko/dvostransko/večslojno/flex)",
  "kolicina": "zahtevana količina",
  "rok_dobave": "zahtevani rok dobave",
  "posebne_zahteve": ["seznam posebnih zahtev (material, certifikati, testi...)"],
  "komponente": [{"naziv": "", "vrednost": "", "ohisje": "", "kolicina": ""}],
  "ocena_zahtevnosti": "Enostavno / Srednje / Zahtevno",
  "opombe": "dodatne opombe",
  "priloge": ["seznam prilog"]
}""",

    "bom_pregled": """Iz spodnje analize/dokumenta izvleci podatke za PREGLED BOM KOSOVNICE.
Vrni SAMO JSON (brez markdown, brez ```), s točno temi ključi:

{
  "naziv_projekta": "naziv projekta/izdelka",
  "verzija_bom": "verzija BOM če je razvidna",
  "datum": "YYYY-MM-DD",
  "stranka": "ime stranke če je razvidno",
  "skupno_komponent": "število unikatnih komponent",
  "skupno_kosov": "skupno število kosov",
  "komponente": [
    {"pozicija": "1", "referenca": "R1,R2", "naziv": "Resistor", "vrednost": "10k", "ohisje": "0402", "kolicina": "2", "dobavitelj": "", "status": "OK/Pozor/Ni na zalogi"}
  ],
  "kriticne_komponente": ["seznam komponent ki zahtevajo pozornost"],
  "manjkajoce_info": ["seznam manjkajočih informacij v BOM"],
  "opombe": "splošne opombe o kakovosti BOM",
  "ocena": "Kompletna / Delno kompletna / Nekompletna"
}""",

    "porocilo": """Iz spodnje analize/dokumenta izvleci podatke za SPLOŠNO POROČILO O PREGLEDU.
Vrni SAMO JSON (brez markdown, brez ```), s točno temi ključi:

{
  "naslov": "naslov poročila",
  "datum": "YYYY-MM-DD",
  "avtor": "",
  "referenca": "referenca projekta/naročila če je razvidna",
  "povzetek": "kratek povzetek ugotovitev (3-5 stavkov)",
  "ugotovitve": [
    {"naslov": "Ugotovitev 1", "opis": "Podroben opis", "resnost": "Info/Opozorilo/Kritično"}
  ],
  "podrobnosti": "podroben opis pregleda",
  "priporocila": ["seznam priporočil"],
  "zakljucek": "zaključek poročila",
  "priloge": ["seznam prilog"]
}""",
}


def extract_pdf_images(pdf_path: str, output_dir: str, start_page: int = 1, min_size: int = 10000) -> list[str]:
    """
    Izvleči slike iz PDF-ja z uporabo PyMuPDF.

    Args:
        pdf_path: pot do PDF datoteke
        output_dir: direktorij za shranjevanje slik
        start_page: od katere strani naprej (0-indexed interno, 1-indexed za uporabnika)
        min_size: minimalna velikost slike v bajtih (filtrira majhne ikone/logotipe)

    Returns:
        seznam poti do ekstrahiranih slik
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF ni nameščen - slike iz PDF ne bodo izvlečene")
        return []

    image_paths = []
    os.makedirs(output_dir, exist_ok=True)

    try:
        pdf_doc = fitz.open(pdf_path)
        for page_num in range(start_page, len(pdf_doc)):
            page = pdf_doc[page_num]
            images = page.get_images(full=True)

            for img_idx, img in enumerate(images):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)

                if not base_image:
                    continue

                image_bytes = base_image["image"]
                # Preskoči majhne slike (logotipi, ikone)
                if len(image_bytes) < min_size:
                    continue

                ext = base_image.get("ext", "png")
                img_filename = f"page{page_num + 1}_img{img_idx + 1}.{ext}"
                img_path = os.path.join(output_dir, img_filename)

                with open(img_path, "wb") as f:
                    f.write(image_bytes)

                image_paths.append(img_path)
                logger.info(f"[PDF] Extracted image: {img_filename} ({len(image_bytes)} bytes)")

        pdf_doc.close()

    except Exception as e:
        logger.error(f"Napaka pri ekstrakciji slik iz PDF: {e}")

    # Če ni slik iz xref, poskusi renderirati strani kot slike
    if not image_paths:
        try:
            pdf_doc = fitz.open(pdf_path)
            for page_num in range(start_page, len(pdf_doc)):
                page = pdf_doc[page_num]
                # Renderiraj stran kot sliko (150 DPI)
                pix = page.get_pixmap(dpi=150)
                img_filename = f"page{page_num + 1}_render.png"
                img_path = os.path.join(output_dir, img_filename)
                pix.save(img_path)
                image_paths.append(img_path)
                logger.info(f"[PDF] Rendered page as image: {img_filename}")

            pdf_doc.close()
        except Exception as e:
            logger.error(f"Napaka pri renderiranju PDF strani: {e}")

    return image_paths


def generate_document(template_type: str, data: dict) -> io.BytesIO:
    """Generiraj Word dokument iz predloge in strukturiranih podatkov."""
    if template_type not in DOCUMENT_TYPES:
        raise ValueError(f"Neznan tip predloge: {template_type}")

    doc = Document()
    _setup_document(doc)

    if template_type == "reklamacija":
        _build_reklamacija(doc, data)
    elif template_type == "rfq_analiza":
        _build_rfq_analiza(doc, data)
    elif template_type == "bom_pregled":
        _build_bom_pregled(doc, data)
    elif template_type == "porocilo":
        _build_porocilo(doc, data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════
# Common helpers
# ═══════════════════════════════════════════════════════════

def _setup_document(doc: Document):
    """Nastavi osnovne stile dokumenta."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _set_cell_shading(cell, color_hex: str):
    """Nastavi ozadje celice."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def _set_cell_text(cell, text: str, bold=False, size=10, color=None, align=None):
    """Nastavi tekst celice z formatiranjem."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_luznar_header(doc: Document):
    """Dodaj Luznar header z imenom in kontaktom."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("LUZNAR")
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.bold = True
    run = p.add_run("  ELECTRONICS")
    run.font.size = Pt(22)
    run.font.color.rgb = GOLD

    p = doc.add_paragraph()
    run = p.add_run("Luznar electronics d.o.o. | Hrastje 52g | SI-4000 Kranj")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run("+386 4 281 88 00 | info@luznar.com")
    run.font.size = Pt(8)
    run.font.color.rgb = GOLD
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Gold separator line
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(8)
    run = p.add_run("━" * 72)
    run.font.size = Pt(6)
    run.font.color.rgb = GOLD


def _add_section_header_table(doc: Document, title: str):
    """Dodaj section header kot tabelo z navy ozadjem (kot v PDF-ju)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, NAVY_HEX)
    _set_cell_text(cell, title, bold=True, size=11, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Set table width to full page
    table.columns[0].width = Cm(16)
    doc.add_paragraph().space_after = Pt(2)


def _add_footer(doc: Document):
    """Dodaj nogo dokumenta."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Luznar electronics d.o.o. | Hrastje 52g, SI-4000 Kranj | info@luznar.com | +386 4 281 88 00")
    run.font.size = Pt(7)
    run.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"www.luznar.com")
    run.font.size = Pt(7)
    run.font.color.rgb = GOLD
    run.bold = True


# ═══════════════════════════════════════════════════════════
# REKLAMACIJA - Supplier Quality Complaint
# Po vzoru 100100306.pdf
# ═══════════════════════════════════════════════════════════

def _build_reklamacija(doc: Document, data: dict):
    """Zgradi Supplier Quality Complaint po vzoru obstoječega PDF-ja."""

    _add_luznar_header(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    run = p.add_run("SUPPLIER QUALITY COMPLAINT")
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.bold = True

    # Complaint Reference
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)
    run = p.add_run("Complaint Reference: ")
    run.font.size = Pt(11)
    run = p.add_run(data.get("complaint_reference", ""))
    run.font.size = Pt(11)
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(16)
    run = p.add_run("Quality Non-Conformance Report")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = GRAY

    # ── SUPPLIER / CUSTOMER table ──
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    _set_cell_shading(table.cell(0, 0), NAVY_HEX)
    _set_cell_shading(table.cell(0, 1), NAVY_HEX)
    _set_cell_shading(table.cell(0, 2), NAVY_HEX)
    _set_cell_shading(table.cell(0, 3), NAVY_HEX)

    # Merge header cells
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 2).merge(table.cell(0, 3))
    _set_cell_text(table.cell(0, 0), "SUPPLIER", bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(0, 2), "CUSTOMER", bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    fields = [
        ("Company:", data.get("supplier_company", ""), "Company:", data.get("customer_company", "LUZNAR electronics d.o.o.")),
        ("Address:", data.get("supplier_address", ""), "Address:", data.get("customer_address", "Hrastje 52g, SI-4000 Kranj, Slovenia")),
        ("Contact:", data.get("supplier_contact", ""), "Contact:", data.get("customer_contact", "info@luznar.com")),
    ]

    for i, (l1, v1, l2, v2) in enumerate(fields):
        row = table.rows[i + 1]
        _set_cell_text(row.cells[0], l1, bold=True, size=9)
        _set_cell_text(row.cells[1], v1, size=9)
        _set_cell_text(row.cells[2], l2, bold=True, size=9)
        _set_cell_text(row.cells[3], v2, size=9)

    doc.add_paragraph().space_after = Pt(10)

    # ── COMPLAINT DETAILS table ──
    _add_section_header_table(doc, "COMPLAINT DETAILS")

    details_data = [
        ("Complaint Date:", data.get("complaint_date", "")),
        ("Article:", data.get("article", "")),
        ("Purchase Order:", data.get("purchase_order", "")),
        ("Delivery Date:", data.get("delivery_date", "")),
        ("Quantity Delivered:", data.get("quantity_delivered", "")),
        ("Quantity Rejected:", data.get("quantity_rejected", "")),
        ("Detection Point:", data.get("detection_point", "")),
    ]

    table = doc.add_table(rows=len(details_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    for i, (label, value) in enumerate(details_data):
        _set_cell_text(table.cell(i, 0), label, bold=True, size=9)
        # Quantity Rejected v rdeči barvi (kot v PDF)
        if label == "Quantity Rejected:" and value:
            _set_cell_text(table.cell(i, 1), value, bold=True, size=9, color=RED)
        elif label == "Article:" and value:
            _set_cell_text(table.cell(i, 1), value, bold=True, size=9)
        else:
            _set_cell_text(table.cell(i, 1), value, size=9)

    doc.add_paragraph().space_after = Pt(10)

    # ── NON-CONFORMANCE DESCRIPTION ──
    _add_section_header_table(doc, "NON-CONFORMANCE DESCRIPTION")

    nc_type = data.get("non_conformance_type", "")
    if nc_type:
        p = doc.add_paragraph()
        p.space_after = Pt(4)
        run = p.add_run("Non-Conformance Type: ")
        run.font.size = Pt(10)
        run.bold = True
        run = p.add_run(nc_type)
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run("Description:")
    run.font.size = Pt(10)
    run.bold = True

    description = data.get("description", "")
    # Split description into paragraphs
    for para_text in description.split("\n"):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph()
            p.space_after = Pt(4)
            run = p.add_run(para_text)
            run.font.size = Pt(10)

    doc.add_paragraph().space_after = Pt(10)

    # ── REQUIRED CORRECTIVE ACTIONS ──
    actions = data.get("corrective_actions", [])
    if actions:
        # New page
        doc.add_page_break()
        _add_luznar_header(doc)

        _add_section_header_table(doc, "REQUIRED CORRECTIVE ACTIONS")

        p = doc.add_paragraph()
        p.space_after = Pt(6)
        run = p.add_run("The following corrective actions are required:")
        run.font.size = Pt(10)
        run.bold = True

        for i, action in enumerate(actions):
            p = doc.add_paragraph()
            p.space_after = Pt(3)
            run = p.add_run(f"{i + 1}. ")
            run.font.size = Pt(10)
            run.bold = True
            run = p.add_run(str(action))
            run.font.size = Pt(10)

    # ── SUPPLIER DECISION & DEADLINES ──
    decision = data.get("supplier_decision", "")
    deadline_4d = data.get("deadline_4d", "")
    deadline_8d = data.get("deadline_8d", "")
    if decision or deadline_4d or deadline_8d:
        doc.add_paragraph().space_after = Pt(10)
        _add_section_header_table(doc, "RESPONSE REQUIREMENTS")

        deadline_data = []
        if decision:
            deadline_data.append(("Supplier Decision:", decision))
        if deadline_4d:
            deadline_data.append(("4D Report Deadline:", deadline_4d))
        if deadline_8d:
            deadline_data.append(("8D Report Deadline:", deadline_8d))

        table = doc.add_table(rows=len(deadline_data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Cm(5)
        table.columns[1].width = Cm(11)

        for i, (label, value) in enumerate(deadline_data):
            _set_cell_text(table.cell(i, 0), label, bold=True, size=9)
            _set_cell_text(table.cell(i, 1), value, bold=True, size=9, color=RED)

    # ── APPENDIX: PHOTOGRAPHIC EVIDENCE ──
    photo_descs = data.get("photo_descriptions", [])
    image_paths = data.get("image_paths", [])

    if photo_descs or image_paths:
        doc.add_page_break()
        _add_luznar_header(doc)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(8)
        p.space_after = Pt(12)
        run = p.add_run("APPENDIX: PHOTOGRAPHIC EVIDENCE")
        run.font.size = Pt(14)
        run.font.color.rgb = NAVY
        run.bold = True

        # Vstavi dejanske slike iz PDF-ja
        if image_paths:
            for i, img_path in enumerate(image_paths):
                try:
                    desc = photo_descs[i] if i < len(photo_descs) else f"Image {i + 1}"
                    p = doc.add_paragraph()
                    p.space_after = Pt(4)
                    run = p.add_run(str(desc))
                    run.font.size = Pt(10)
                    run.bold = True

                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.space_after = Pt(12)
                    run = p.add_run()
                    run.add_picture(img_path, width=Cm(14))
                except Exception:
                    p = doc.add_paragraph()
                    p.space_after = Pt(12)
                    run = p.add_run(f"[Napaka pri vstavljanju slike: {img_path}]")
                    run.font.size = Pt(9)
                    run.font.color.rgb = GRAY
                    run.italic = True
        else:
            # Samo opisi brez slik
            for desc in photo_descs:
                p = doc.add_paragraph()
                p.space_after = Pt(6)
                run = p.add_run(str(desc))
                run.font.size = Pt(10)
                run.bold = True

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_after = Pt(12)
                run = p.add_run("[Fotografija - vstavi ročno]")
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY
                run.italic = True

    _add_footer(doc)


# ═══════════════════════════════════════════════════════════
# RFQ ANALIZA
# ═══════════════════════════════════════════════════════════

def _build_rfq_analiza(doc: Document, data: dict):
    """Zgradi analizo RFQ povpraševanja."""
    _add_luznar_header(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    p.space_after = Pt(16)
    run = p.add_run("ANALIZA POVPRAŠEVANJA (RFQ)")
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.bold = True

    # ── Podatki o povpraševanju ──
    _add_section_header_table(doc, "PODATKI O POVPRAŠEVANJU")

    fields = [
        ("Št. RFQ:", data.get("stevilka_rfq", "")),
        ("Datum prejema:", data.get("datum_prejema", "")),
        ("Stranka:", data.get("stranka", "")),
        ("Kontakt:", data.get("kontakt", "")),
        ("Količina:", data.get("kolicina", "")),
        ("Rok dobave:", data.get("rok_dobave", "")),
        ("Tip vezja:", data.get("tip_vezja", "")),
        ("Ocena zahtevnosti:", data.get("ocena_zahtevnosti", "")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    for i, (label, value) in enumerate(fields):
        _set_cell_text(table.cell(i, 0), label, bold=True, size=9)
        _set_cell_text(table.cell(i, 1), value, size=9)

    doc.add_paragraph().space_after = Pt(10)

    # ── Opis projekta ──
    opis = data.get("opis_projekta", "")
    if opis:
        _add_section_header_table(doc, "OPIS PROJEKTA")
        p = doc.add_paragraph()
        p.space_after = Pt(4)
        run = p.add_run(opis)
        run.font.size = Pt(10)
        doc.add_paragraph().space_after = Pt(10)

    # ── Posebne zahteve ──
    zahteve = data.get("posebne_zahteve", [])
    if zahteve:
        _add_section_header_table(doc, "POSEBNE ZAHTEVE")
        for z in zahteve:
            p = doc.add_paragraph(style="List Bullet")
            p.space_after = Pt(2)
            run = p.add_run(str(z))
            run.font.size = Pt(10)
        doc.add_paragraph().space_after = Pt(10)

    # ── Komponente ──
    komponente = data.get("komponente", [])
    if komponente:
        _add_section_header_table(doc, "SEZNAM KOMPONENT")
        headers = ["Naziv", "Vrednost", "Ohišje", "Količina"]
        table = doc.add_table(rows=1 + len(komponente), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            _set_cell_shading(table.cell(0, i), LIGHT_GRAY_HEX)
            _set_cell_text(table.cell(0, i), h, bold=True, size=9, color=NAVY)

        for r, k in enumerate(komponente):
            if isinstance(k, dict):
                _set_cell_text(table.cell(r + 1, 0), k.get("naziv", ""), size=9)
                _set_cell_text(table.cell(r + 1, 1), k.get("vrednost", ""), size=9)
                _set_cell_text(table.cell(r + 1, 2), k.get("ohisje", ""), size=9)
                _set_cell_text(table.cell(r + 1, 3), k.get("kolicina", ""), size=9)

    # ── Opombe ──
    opombe = data.get("opombe", "")
    if opombe:
        doc.add_paragraph().space_after = Pt(10)
        _add_section_header_table(doc, "OPOMBE")
        p = doc.add_paragraph()
        run = p.add_run(opombe)
        run.font.size = Pt(10)

    _add_footer(doc)


# ═══════════════════════════════════════════════════════════
# BOM PREGLED
# ═══════════════════════════════════════════════════════════

def _build_bom_pregled(doc: Document, data: dict):
    """Zgradi pregled BOM kosovnice."""
    _add_luznar_header(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    p.space_after = Pt(16)
    run = p.add_run("PREGLED KOSOVNICE (BOM)")
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.bold = True

    # ── Podatki ──
    _add_section_header_table(doc, "PODATKI O KOSOVNICI")

    fields = [
        ("Projekt:", data.get("naziv_projekta", "")),
        ("Verzija BOM:", data.get("verzija_bom", "")),
        ("Datum:", data.get("datum", "")),
        ("Stranka:", data.get("stranka", "")),
        ("Skupno komponent:", data.get("skupno_komponent", "")),
        ("Skupno kosov:", data.get("skupno_kosov", "")),
        ("Ocena:", data.get("ocena", "")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    for i, (label, value) in enumerate(fields):
        _set_cell_text(table.cell(i, 0), label, bold=True, size=9)
        _set_cell_text(table.cell(i, 1), value, size=9)

    doc.add_paragraph().space_after = Pt(10)

    # ── Komponente tabela ──
    komponente = data.get("komponente", [])
    if komponente:
        _add_section_header_table(doc, "SEZNAM KOMPONENT")
        headers = ["#", "Ref.", "Naziv", "Vrednost", "Ohišje", "Kol.", "Status"]
        table = doc.add_table(rows=1 + len(komponente), cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            _set_cell_shading(table.cell(0, i), LIGHT_GRAY_HEX)
            _set_cell_text(table.cell(0, i), h, bold=True, size=8, color=NAVY)

        for r, k in enumerate(komponente):
            if isinstance(k, dict):
                _set_cell_text(table.cell(r + 1, 0), k.get("pozicija", ""), size=8)
                _set_cell_text(table.cell(r + 1, 1), k.get("referenca", ""), size=8)
                _set_cell_text(table.cell(r + 1, 2), k.get("naziv", ""), size=8)
                _set_cell_text(table.cell(r + 1, 3), k.get("vrednost", ""), size=8)
                _set_cell_text(table.cell(r + 1, 4), k.get("ohisje", ""), size=8)
                _set_cell_text(table.cell(r + 1, 5), str(k.get("kolicina", "")), size=8)
                status = k.get("status", "")
                color = RED if status and "ni" in status.lower() else None
                _set_cell_text(table.cell(r + 1, 6), status, size=8, color=color)

        doc.add_paragraph().space_after = Pt(10)

    # ── Kritične komponente ──
    kriticne = data.get("kriticne_komponente", [])
    if kriticne:
        _add_section_header_table(doc, "KRITIČNE KOMPONENTE")
        for k in kriticne:
            p = doc.add_paragraph(style="List Bullet")
            p.space_after = Pt(2)
            run = p.add_run(str(k))
            run.font.size = Pt(10)
        doc.add_paragraph().space_after = Pt(10)

    # ── Manjkajoče info ──
    manjka = data.get("manjkajoce_info", [])
    if manjka:
        _add_section_header_table(doc, "MANJKAJOČE INFORMACIJE")
        for m in manjka:
            p = doc.add_paragraph(style="List Bullet")
            p.space_after = Pt(2)
            run = p.add_run(str(m))
            run.font.size = Pt(10)

    _add_footer(doc)


# ═══════════════════════════════════════════════════════════
# POROČILO O PREGLEDU
# ═══════════════════════════════════════════════════════════

def _build_porocilo(doc: Document, data: dict):
    """Zgradi splošno poročilo o pregledu."""
    _add_luznar_header(doc)

    naslov = data.get("naslov", "Poročilo o pregledu")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    p.space_after = Pt(16)
    run = p.add_run(naslov.upper())
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.bold = True

    # ── Podatki ──
    _add_section_header_table(doc, "PODATKI")

    fields = [
        ("Datum:", data.get("datum", "")),
        ("Avtor:", data.get("avtor", "")),
        ("Referenca:", data.get("referenca", "")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)

    for i, (label, value) in enumerate(fields):
        _set_cell_text(table.cell(i, 0), label, bold=True, size=9)
        _set_cell_text(table.cell(i, 1), value, size=9)

    doc.add_paragraph().space_after = Pt(10)

    # ── Povzetek ──
    povzetek = data.get("povzetek", "")
    if povzetek:
        _add_section_header_table(doc, "POVZETEK")
        p = doc.add_paragraph()
        p.space_after = Pt(4)
        run = p.add_run(povzetek)
        run.font.size = Pt(10)
        doc.add_paragraph().space_after = Pt(10)

    # ── Ugotovitve ──
    ugotovitve = data.get("ugotovitve", [])
    if ugotovitve:
        _add_section_header_table(doc, "UGOTOVITVE")
        headers = ["Ugotovitev", "Opis", "Resnost"]
        table = doc.add_table(rows=1 + len(ugotovitve), cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            _set_cell_shading(table.cell(0, i), LIGHT_GRAY_HEX)
            _set_cell_text(table.cell(0, i), h, bold=True, size=9, color=NAVY)

        for r, u in enumerate(ugotovitve):
            if isinstance(u, dict):
                _set_cell_text(table.cell(r + 1, 0), u.get("naslov", ""), bold=True, size=9)
                _set_cell_text(table.cell(r + 1, 1), u.get("opis", ""), size=9)
                resnost = u.get("resnost", "")
                color = RED if "kritič" in resnost.lower() else None
                _set_cell_text(table.cell(r + 1, 2), resnost, size=9, color=color)

        doc.add_paragraph().space_after = Pt(10)

    # ── Podrobnosti ──
    podrobnosti = data.get("podrobnosti", "")
    if podrobnosti:
        _add_section_header_table(doc, "PODROBNOSTI PREGLEDA")
        p = doc.add_paragraph()
        run = p.add_run(podrobnosti)
        run.font.size = Pt(10)
        doc.add_paragraph().space_after = Pt(10)

    # ── Priporočila ──
    priporocila = data.get("priporocila", [])
    if priporocila:
        _add_section_header_table(doc, "PRIPOROČILA")
        for i, pr in enumerate(priporocila):
            p = doc.add_paragraph()
            p.space_after = Pt(3)
            run = p.add_run(f"{i + 1}. ")
            run.font.size = Pt(10)
            run.bold = True
            run = p.add_run(str(pr))
            run.font.size = Pt(10)
        doc.add_paragraph().space_after = Pt(10)

    # ── Zaključek ──
    zakljucek = data.get("zakljucek", "")
    if zakljucek:
        _add_section_header_table(doc, "ZAKLJUČEK")
        p = doc.add_paragraph()
        run = p.add_run(zakljucek)
        run.font.size = Pt(10)

    _add_footer(doc)

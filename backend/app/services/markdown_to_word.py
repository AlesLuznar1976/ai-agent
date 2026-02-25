"""
Markdown to Word - Pretvori markdown tekst v Word (.docx) dokument.

Podpira: naslove (H1-H3), bold, italic, tabele, sezname, code bloke.
"""

import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def markdown_to_docx(markdown_text: str, title: str = "Analiza") -> io.BytesIO:
    """Pretvori markdown v Word dokument, vrne BytesIO."""
    doc = Document()

    # Stil dokumenta
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Naslov dokumenta
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Tabela - zaznaj po | na začetku
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Heading
        if line.startswith("### "):
            doc.add_heading(_clean_inline(line[4:]), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(_clean_inline(line[3:]), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(_clean_inline(line[2:]), level=1)
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Bullet list
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, text)
            i += 1
            continue

        # Numbered list
        numbered = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if numbered:
            text = numbered.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, text)
            i += 1
            continue

        # Prazna vrstica
        if not line.strip():
            i += 1
            continue

        # Navaden odstavek
        p = doc.add_paragraph()
        _add_rich_text(p, line)
        i += 1

    # Shrani v BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _clean_inline(text: str) -> str:
    """Odstrani markdown inline formatiranje."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def _add_rich_text(paragraph, text: str):
    """Doda tekst z bold/italic/code podporo v odstavek."""
    # Parse inline markdown: **bold**, *italic*, `code`
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))"
    for match in re.finditer(pattern, text):
        full = match.group(0)
        if full.startswith("**") and full.endswith("**"):
            run = paragraph.add_run(full[2:-2])
            run.bold = True
        elif full.startswith("*") and full.endswith("*"):
            run = paragraph.add_run(full[1:-1])
            run.italic = True
        elif full.startswith("`") and full.endswith("`"):
            run = paragraph.add_run(full[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(full)


def _add_table(doc: Document, table_lines: list[str]):
    """Pretvori markdown tabelo v Word tabelo."""
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Preskoči separator vrstico (---|----|---)
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = _clean_inline(cell_text)
                # Header row bold
                if r_idx == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(10)

    doc.add_paragraph()  # Razmik po tabeli

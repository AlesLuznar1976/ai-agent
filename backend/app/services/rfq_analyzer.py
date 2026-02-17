"""
RFQ Deep Analysis - poglobljena analiza povpraševanj.

Ko sistem zazna email kategorije RFQ, Naročilo ali Ponudba,
ta modul:
1. Prenese in parsira vse priloge (Excel, Word, PDF, ZIP)
2. Z LLM (Claude) izvleči strukturirane informacije
3. Ustvari strukturiran seznam: kaj je stranka podala in kaj manjka
"""

import io
import json
import re
import zipfile
from typing import Optional

from sqlalchemy.orm import Session

from app.crud import emaili as crud_emaili
from app.llm import get_llm_router, TaskType
from app.models.email import RfqPodkategorija
from app.services.attachment_processor import (
    download_attachment,
    extract_pdf_text,
    fetch_attachment_metadata,
)
from app.services.email_sync import get_ms_graph_token


# ============================================================
# Parsiranje datotek
# ============================================================

def extract_excel_text(content_bytes: bytes, filename: str = "") -> str:
    """Parsira Excel datoteko in vrne tekst vseh listov/celic."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(filename=io.BytesIO(content_bytes), read_only=True, data_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- List: {sheet_name} ---")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                # Preskoči prazne vrstice
                if any(cells):
                    parts.append("\t".join(cells))
                row_count += 1
                if row_count >= 500:  # Omejitev vrstic
                    parts.append(f"... (prekinjen po {row_count} vrsticah)")
                    break

        wb.close()
        text = "\n".join(parts)
        return text[:30000]  # Omejitev znakov

    except Exception as e:
        print(f"Excel parsiranje napaka ({filename}): {e}")
        return f"[Napaka pri parsiranju Excel: {e}]"


def extract_word_text(content_bytes: bytes, filename: str = "") -> str:
    """Parsira Word (.docx) datoteko - odstavki in tabele."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content_bytes))
        parts = []

        # Odstavki
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tabele
        for i, table in enumerate(doc.tables):
            parts.append(f"\n--- Tabela {i + 1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))

        text = "\n".join(parts)
        return text[:30000]

    except Exception as e:
        print(f"Word parsiranje napaka ({filename}): {e}")
        return f"[Napaka pri parsiranju Word: {e}]"


def parse_attachment(content_bytes: bytes, filename: str) -> str:
    """Dispatcher: glede na extension kliče pravo funkcijo za parsiranje."""
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        return extract_pdf_text(content_bytes)
    elif name_lower.endswith((".xlsx", ".xls")):
        return extract_excel_text(content_bytes, filename)
    elif name_lower.endswith((".docx", ".doc")):
        return extract_word_text(content_bytes, filename)
    elif name_lower.endswith((".txt", ".csv")):
        try:
            return content_bytes.decode("utf-8", errors="replace")[:30000]
        except Exception:
            return "[Napaka pri branju tekstovne datoteke]"
    else:
        return ""


async def extract_zip_contents(
    content_bytes: bytes,
    token: str,
    message_id: str = "",
    mailbox: str = "",
) -> list[dict]:
    """Razpakira ZIP in rekurzivno parsira vsebino (Excel, Word, PDF v ZIP-u).

    Returns:
        list[dict] z {"filename": ..., "text": ...} za vsako datoteko v ZIP-u
    """
    results = []

    try:
        with zipfile.ZipFile(io.BytesIO(content_bytes)) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue

                name = info.filename
                name_lower = name.lower()

                # Preskoči nepomembne datoteke
                if name_lower.startswith("__macosx") or name_lower.startswith("."):
                    continue

                # Parsira samo znane tipe
                parseable_extensions = (
                    ".pdf", ".xlsx", ".xls", ".docx", ".doc",
                    ".txt", ".csv",
                )
                if not any(name_lower.endswith(ext) for ext in parseable_extensions):
                    results.append({"filename": name, "text": f"[Datoteka: {name} - ni parsirana]"})
                    continue

                try:
                    inner_bytes = zf.read(info)
                    text = parse_attachment(inner_bytes, name)
                    if text:
                        results.append({"filename": name, "text": text})
                except Exception as e:
                    results.append({"filename": name, "text": f"[Napaka pri parsiranju: {e}]"})

    except zipfile.BadZipFile:
        results.append({"filename": "ZIP", "text": "[Neveljaven ZIP arhiv]"})
    except Exception as e:
        results.append({"filename": "ZIP", "text": f"[Napaka pri razpakiranju: {e}]"})

    return results


# ============================================================
# LLM analiza
# ============================================================

def _build_analysis_prompt(email_body: str, attachment_texts: list[dict]) -> str:
    """Sestavi PCB/SMT domenski prompt za analizo RFQ emaila."""

    # Sestavi attachment besedila
    att_sections = []
    for att in attachment_texts:
        att_sections.append(
            f"=== Priloga: {att['filename']} ===\n{att['text'][:8000]}"
        )
    attachments_text = "\n\n".join(att_sections) if att_sections else "Ni prilog ali ni mogoče parsirati prilog."

    return f"""Si strokovni AI asistent za podjetje Luznar Electronics d.o.o. (izdelava PCB, SMT sestav, elektronskih sklopov).

Analiziraj naslednje povpraševanje (RFQ email) in iz njega izvleci VSE relevantne informacije za pripravo ponudbe.

**Email vsebina:**
{email_body[:4000]}

**Vsebina prilog:**
{attachments_text}

**Naloga:**
1. Identificiraj stranko (ime, kontakt, email)
2. Določi tip povpraševanja (PCB / SMT / Sestav / Drugo)
3. Izvleci vse izdelke s specifikacijami (material, sloji, dimenzije, količina, površinska obdelava, ...)
4. Identificiraj katere dokumente je stranka priložila (BOM, Gerber, risbe, specifikacije)
5. Določi katere informacije je stranka PODALA
6. Določi katere informacije MANJKAJO za pripravo ponudbe
7. Napiši kratek povzetek
8. Oceni prioriteto (Visoka/Srednja/Nizka)
9. Predlagaj naslednje korake
10. Določi RFQ pod-kategorijo:
    - Kompletno: Ima BOM + Gerber + specifikacije + količino
    - Nepopolno: Ima nekatere dokumente, a ne vseh
    - Povpraševanje: Splošno vprašanje brez tehničnih dokumentov
    - Repeat Order: Ponovitev prejšnjega naročila (ključne besede: ponovitev, repeat, reorder, enako kot)

**PCB specifikacije ki jih tipično potrebujemo:**
- Material (FR4, Rogers, CEM, ...)
- Število slojev
- Debelina PCB
- Debelina bakra
- Površinska obdelava (HASL, ENIG, OSP, ...)
- Barva spajkalne maske
- Dimenzije plošče
- Količina
- Tolerances (impedanca, ...)
- Posebne zahteve (flex, rigid-flex, HDI, ...)

**SMT specifikacije ki jih tipično potrebujemo:**
- BOM (Bill of Materials)
- Gerber datoteke
- Pick & place datoteka
- Količina
- Testne zahteve (ICT, funkcionalni test, ...)
- Posebne zahteve (conformal coating, potting, ...)

POMEMBNO: Vrni IZKLJUČNO EN SAM veljaven JSON objekt. Brez razlage, brez markdown, brez komentarjev, brez oštevilčenja - SAMO JSON:
{{
    "stranka": {{
        "ime": "ime podjetja ali osebe",
        "kontakt": "kontaktna oseba",
        "email": "email naslov"
    }},
    "tip_povprasevanja": "PCB / SMT / Sestav / Drugo",
    "izdelki": [
        {{
            "naziv": "ime izdelka ali opis",
            "kolicina": null,
            "specifikacije": {{}}
        }}
    ],
    "prilozeni_dokumenti": [
        {{
            "ime": "ime datoteke",
            "tip": "BOM / Gerber / Specifikacija / Risba / Drugo",
            "vsebina_povzetek": "kratek opis vsebine"
        }}
    ],
    "podano_od_stranke": ["seznam informacij ki jih je stranka podala"],
    "manjkajoci_podatki": ["seznam informacij ki manjkajo za ponudbo"],
    "povzetek": "kratek povzetek povpraševanja",
    "prioriteta": "Visoka/Srednja/Nizka",
    "rfq_podkategorija": "Kompletno/Nepopolno/Povpraševanje/Repeat Order",
    "priporoceni_naslednji_koraki": ["seznam priporočenih korakov"]
}}"""


async def _run_llm_analysis(prompt: str) -> dict:
    """Pošlje prompt na Claude preko LLM routerja in parsira JSON odgovor."""
    llm = get_llm_router()

    # Začasno povečaj timeout za lokalni LLM (velike priloge)
    original_timeout = llm.local_llm.timeout
    llm.local_llm.timeout = 300.0  # 5 minut za analizo z prilogami

    try:
        response = await llm.complete(
            prompt,
            task_type=TaskType.COMPLEX_REASONING,
            force_local=True,
        )
    finally:
        llm.local_llm.timeout = original_timeout

    # Parse JSON iz odgovora
    text = response.strip()

    # Poskusi direktno
    try:
        return json.loads(text)
    except (json.JSONDecodeError, ValueError):
        pass

    # Poišči JSON v code blockih
    code_block = re.search(r"```(?:json)?\s*\n?(.*?)\n?\s*```", text, re.DOTALL)
    if code_block:
        try:
            return json.loads(code_block.group(1).strip())
        except (json.JSONDecodeError, ValueError):
            pass

    # Poišči največji {..} blok ki se parsira kot JSON
    # (lokalni LLM včasih vrne več manjših JSON blokov namesto enega)
    brace_starts = [m.start() for m in re.finditer(r"\{", text)]
    for start in brace_starts:
        # Najdi ujemajoč zaključni }
        depth = 0
        for i in range(start, len(text)):
            if text[i] == "{":
                depth += 1
            elif text[i] == "}":
                depth -= 1
                if depth == 0:
                    candidate = text[start:i + 1]
                    try:
                        parsed = json.loads(candidate)
                        # Preveri da ima vsaj 3 ključe (ni fragmentiran)
                        if isinstance(parsed, dict) and len(parsed) >= 3:
                            return parsed
                    except (json.JSONDecodeError, ValueError):
                        pass
                    break

    # Fallback: poskusi katerikoli {..} blok
    brace_match = re.search(r"\{.*\}", text, re.DOTALL)
    if brace_match:
        try:
            return json.loads(brace_match.group(0))
        except (json.JSONDecodeError, ValueError):
            pass

    raise ValueError(f"Ni mogoče izvleči JSON iz LLM odgovora: {text[:300]}")


# ============================================================
# RFQ Pod-kategorizacija (Faza 2 - deterministična)
# ============================================================

def _determine_subcategory(result: dict, email_body: str) -> RfqPodkategorija:
    """Deterministično določi RFQ pod-kategorijo iz parsiranih prilog in vsebine.

    Faza 2 override: po parsiranju prilog imamo natančnejše podatke kot
    v Fazi 1 (ki vidi le imena prilog).
    """
    body_lower = (email_body or "").lower()

    # Repeat order ključne besede
    repeat_keywords = [
        "ponovitev", "repeat", "reorder", "re-order",
        "enako kot", "isto kot", "ponovi naročilo", "same as before",
    ]
    if any(kw in body_lower for kw in repeat_keywords):
        return RfqPodkategorija.REPEAT_ORDER

    # Preveri parsirane priloge
    dokumenti = result.get("prilozeni_dokumenti") or []
    has_bom = any(
        (d.get("tip") or "").upper() == "BOM"
        for d in dokumenti if isinstance(d, dict)
    )
    has_gerber = any(
        (d.get("tip") or "").upper() == "GERBER"
        for d in dokumenti if isinstance(d, dict)
    )
    has_spec = any(
        (d.get("tip") or "").upper() in ("SPECIFIKACIJA", "RISBA")
        for d in dokumenti if isinstance(d, dict)
    )

    # Preveri količino v izdelkih
    izdelki = result.get("izdelki") or []
    has_quantity = any(
        i.get("kolicina") is not None
        for i in izdelki if isinstance(i, dict)
    )

    if has_bom and has_gerber and has_spec and has_quantity:
        return RfqPodkategorija.KOMPLETNO
    elif has_bom or has_gerber or has_spec:
        return RfqPodkategorija.NEPOPOLNO
    else:
        return RfqPodkategorija.POVPRASEVANJE


# ============================================================
# Glavni pipeline
# ============================================================

async def analyze_rfq_email(db: Session, email_id: int) -> dict:
    """Celoten pipeline za analizo enega RFQ emaila.

    1. Naloži email iz DB
    2. Pridobi MS Graph token
    3. Prenesi in parsiraj vse priloge (vključno z ZIP vsebino)
    4. Sestavi prompt z email telom + vsemi parsiranimi prilogami
    5. Pošlji na Claude za analizo
    6. Shrani rezultat v analiza_rezultat, nastavi analiza_status="Končano"
    7. Vrne strukturiran JSON
    """
    db_email = crud_emaili.get_email_by_id(db, email_id)
    if not db_email:
        raise ValueError(f"Email {email_id} ne obstaja")

    # Nastavi status "V obdelavi"
    crud_emaili.update_email(db, email_id, analiza_status="V obdelavi")

    try:
        # Pridobi token
        token = await get_ms_graph_token()

        # Ugotovi mailbox iz izvlečenih podatkov
        mailbox = None
        if db_email.izvleceni_podatki:
            try:
                izvl = json.loads(db_email.izvleceni_podatki)
                mailbox = izvl.get("mailbox")
            except (json.JSONDecodeError, TypeError):
                pass

        # Prenesi in parsiraj priloge
        attachment_texts = []
        if token and db_email.priloge:
            try:
                priloge = json.loads(db_email.priloge)
            except (json.JSONDecodeError, TypeError):
                priloge = []

            if priloge:
                # Pridobi metadata iz Graph za ID-je
                att_metadata = await fetch_attachment_metadata(
                    token, db_email.outlook_id, mailbox=mailbox
                )

                for att_meta in att_metadata:
                    att_id = att_meta.get("id", "")
                    att_name = att_meta.get("name", "unknown")

                    # Prenesi vsebino
                    content = await download_attachment(
                        token, db_email.outlook_id, att_id, mailbox=mailbox
                    )
                    if not content:
                        attachment_texts.append({
                            "filename": att_name,
                            "text": "[Prenos ni uspel]",
                        })
                        continue

                    # ZIP: razpakira in parsira vsebino
                    if att_name.lower().endswith(".zip"):
                        zip_contents = await extract_zip_contents(
                            content, token, db_email.outlook_id, mailbox or ""
                        )
                        for zc in zip_contents:
                            attachment_texts.append({
                                "filename": f"{att_name}/{zc['filename']}",
                                "text": zc["text"],
                            })
                    else:
                        # Parsira direktno
                        text = parse_attachment(content, att_name)
                        if text:
                            attachment_texts.append({
                                "filename": att_name,
                                "text": text,
                            })
                        else:
                            attachment_texts.append({
                                "filename": att_name,
                                "text": f"[Datoteka tipa {att_name.split('.')[-1]} - ni parsirana]",
                            })

        # Sestavi prompt in poženi LLM analizo
        email_body = db_email.telo or ""
        prompt = _build_analysis_prompt(email_body, attachment_texts)
        result = await _run_llm_analysis(prompt)

        # Faza 2: Deterministična pod-kategorija (override LLM)
        rfq_podkat = _determine_subcategory(result, email_body)
        result["rfq_podkategorija"] = rfq_podkat.value

        # Shrani rezultat + pod-kategorijo
        crud_emaili.update_email(
            db, email_id,
            analiza_status="Končano",
            analiza_rezultat=result,
            rfq_podkategorija=rfq_podkat.value,
        )

        return result

    except Exception as e:
        # Shrani napako
        error_msg = str(e)
        print(f"RFQ analysis error for email {email_id}: {error_msg}")
        crud_emaili.update_email(
            db, email_id,
            analiza_status="Napaka",
            analiza_rezultat={"error": error_msg},
        )
        raise


# ============================================================
# Batch processor
# ============================================================

async def process_pending_analyses(db: Session) -> int:
    """Obdela vse emaile z analiza_status='Čaka'.

    Returns:
        Število uspešno obdelanih emailov
    """
    pending = crud_emaili.list_emails_pending_analysis(db)
    if not pending:
        return 0

    analyzed = 0
    for db_email in pending:
        try:
            await analyze_rfq_email(db, db_email.id)
            analyzed += 1
        except Exception as e:
            print(f"RFQ batch analysis error for email {db_email.id}: {e}")
            # Status je že nastavljen na "Napaka" v analyze_rfq_email

    return analyzed
